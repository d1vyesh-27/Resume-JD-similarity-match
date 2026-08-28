"""
Document parsing module for extracting text from PDFs and DOCX files.

IMPORTANT IMPLEMENTATION CONSTRAINT:
When implementing components in this module, first explain in comments:
1. What the concept is.
2. Why it is needed.
3. What the input is.
4. What the output is.
5. What the relevant library function does.
Do not dump unexplained code.
"""

import pdfplumber
import docx

def extract_pdf_text(file) -> str:
    """
    1. Concept: PDF Text Extraction
    2. Why: To convert uploaded PDF resumes into raw text strings.
    3. Input: A file-like object (e.g., BytesIO from Streamlit).
    4. Output: Extracted text as a single string.
    5. Library: pdfplumber.open() reads the PDF, page.extract_text() gets text.
    """
    with pdfplumber.open(file) as pdf:
        text = []
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:  # Check if text was actually extracted from the page
                text.append(page_text)
    return " ".join(text)

def extract_docx_text(file) -> str:
    """
    1. Concept: DOCX Text Extraction
    2. Why: To convert uploaded Word documents into raw text strings.
    3. Input: A file-like object (e.g., BytesIO from Streamlit).
    4. Output: Extracted text as a single string.
    5. Library: docx.Document() reads the DOCX, doc.paragraphs gets text blocks.
    """
    doc = docx.Document(file)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return " ".join(paragraphs)

def extract_resume_text(file, filename: str) -> str:
    """
    1. Concept: File Router
    2. Why: To automatically route the uploaded file to the correct parser based on extension.
    3. Input: File-like object and the string filename.
    4. Output: Extracted text as a single string.
    5. Library: Standard Python string methods (.lower(), .endswith()).
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return extract_pdf_text(file)
    elif filename_lower.endswith(".docx"):
        return extract_docx_text(file)
    else:
        raise ValueError(f"Unsupported file format. Please upload a PDF or DOCX file. File: {filename}")
